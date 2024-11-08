from docx import Document
import pandas as pd

# Caminho do arquivo Excel
excel_path = 'C:\\Users\\BurgerKing\\Documents\\EQUIPAMENTOS_REPARADOS.xlsx'

# Ler a planilha Excel ignorando a primeira linha (cabeçalho)
df = pd.read_excel(excel_path, sheet_name='Planilha1')

# Selecionar apenas as primeiras 5 linhas de dados
#dados = df.iloc[:21]

# Criar um novo documento Word
doc = Document()

# Adicionar uma tabela ao documento, com o mesmo número de linhas e colunas dos dados
#table = doc.add_table(rows=20, cols=len(dados.columns))
table = doc.add_table(rows=20, cols=4)

# Preencher a tabela com os dados da planilha
for i, row in enumerate(df.values):
    for j, cell in enumerate(row):
        table.cell(i, j).text = str(cell)  # Converte o valor para string antes de inserir


# Save the document
output_path = "C:\\Users\\BurgerKing\\Desktop\\ProjetoA\\Document_with_Tables.docx"
doc.save(output_path)


print("Dados da planilha foram inseridos no documento Word.")
