from docx import Document
import pandas as pd
import math

# Caminho do arquivo Excel
excel_path = 'C:\\Users\\BurgerKing\\Downloads\\LAUDO - SUCATA BK.xlsx'

# Ler a planilha do Excel
df = pd.read_excel(excel_path, sheet_name='LAUDO 113')

# Número de linhas por documento
linhas_por_documento = 20

# Calcular quantos documentos são necessários (divisão arredondada para cima)
num_documentos = math.ceil(df.shape[0] / linhas_por_documento)

# Criar documentos dinamicamente com base no número de documentos necessários
for i in range(num_documentos):
    # Obter as linhas do DataFrame para o documento atual
    inicio = i * linhas_por_documento
    fim = inicio + linhas_por_documento
    dados_documento = df.iloc[inicio:fim]

    # Criar um novo documento Word
    doc = Document()

    # Criar a tabela principal de itens (linhas + 2 colunas adicionais)
    table = doc.add_table(rows=linhas_por_documento, cols=len(dados_documento.columns) + 2)

    # Preencher a tabela principal com os dados e adicionar a coluna "Item" e uma coluna vazia
    for j in range(linhas_por_documento):
        # Coluna "Item" com numeração de 1 a 20
        table.cell(j, 0).text = str(j + 1)
        
        # Coluna vazia
        table.cell(j, 1).text = ""
        
        # Preencher o restante das colunas com os dados do Excel, se disponíveis
        if j < len(dados_documento):
            for k, cell in enumerate(dados_documento.iloc[j]):
                table.cell(j, k + 2).text = str(cell)
        else:
            # Preencher células vazias se não houver mais dados
            for k in range(len(dados_documento.columns)):
                table.cell(j, k + 2).text = ""

    

    # Adicionar uma tabela 1x2 abaixo de cada linha da tabela principal
    for _ in range(linhas_por_documento + 1):
        # Criar a tabela 1x2
        doc.add_paragraph()
        sub_table = doc.add_table(rows=1, cols=2)
        # Deixar as duas células vazias ou adicionar um preenchimento específico se necessário
        sub_table.cell(0, 0).text = ""
        sub_table.cell(0, 1).text = ""
        # Adicionar uma linha em branco após cada tabela 1x2 para espaçamento visual
        

    # Salvar o documento com um nome único
    doc_path = f'C:\\Users\\BurgerKing\\Documents\\Equipamentos_Reparados_Parte{i+1}.docx'
    doc.save(doc_path)
    
    print(f"Documento {i+1} criado com as linhas de {inicio+1} a {fim}.")
