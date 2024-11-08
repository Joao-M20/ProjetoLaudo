
from docx import Document
from docx.shared import Inches

# Criar um novo documento
doc = Document()

# Quantidade de TAGs e EQUIPAMENTOS para preencher
num_linhas = 20  # Ajuste conforme necessário

for i in range(num_linhas):
    # Criar uma tabela com 1 linha e 2 colunas
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3.0)  # Define largura para a coluna "TAG"
    table.columns[0].height = Inches(2.5)
    table.columns[1].width = Inches(4.0)  # Define largura para a coluna "EQUIPAMENTO"
    table.columns[1].height = Inches(2.5)

    # Preencher as células com "TAG" e "EQUIPAMENTO"
    table.cell(0, 0).text = f"TAG {i + 1}"
    table.cell(0, 1).text = "EQUIPAMENTO"

    # Espaço após cada tabela
    doc.add_paragraph()  # Adiciona um parágrafo em branco para espaçamento

# Salvar o documento
output_path = "C:\\Users\\BurgerKing\\Documents\\Documento_Gerado.docx"
doc.save(output_path)

print(f"Documento criado e salvo em: {output_path}")
