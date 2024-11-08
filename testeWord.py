from docx import Document
import pandas as pd

# Carregar o documento criado
doc = Document("Document_with_Tables.docx")

#Carregar arquivo excel
excel_path = 'C:\\Users\\BurgerKing\\Documents\\EQUIPAMENTOS_REPARADOS.xlsx'

df =pd.read_excel(excel_path, sheet_name='Planilha1')



# Acessar a primeira tabela (6x5) e preencher com um exemplo
table_6x5 = doc.tables[0]
for row in range(6):
    for col in range(5):
        table_6x5.cell(row, col).text = f"Linha {row+1}, Coluna {col+1}"

# Acessar a segunda tabela (2x2) e preencher com um exemplo
table_2x2 = doc.tables[1]
for row in range(2):
    for col in range(2):
        table_2x2.cell(row, col).text = f"Célula {row+1}, {col+1}"



# Salvar o documento com as tabelas preenchidas
doc.save("Document_with_Tables_Filled.docx")
