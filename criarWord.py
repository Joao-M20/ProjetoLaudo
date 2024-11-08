from docx import Document
import pandas as pd

excel_path = 'C:\\Users\\BurgerKing\\Documents\\EQUIPAMENTOS_REPARADOS.xlsx'

df = pd.read_excel(excel_path, sheet_name='Planilha1', skiprows=1)

dados = df.iloc[:5]

# Create a new Word Document
doc = Document()

doc.add_paragraph()
doc.add_paragraph()

# Add the first table (6x5) to the document
table_6x5 = doc.add_table(rows=6, cols=5)

# Add spacing between tables
doc.add_paragraph()  

# Add the second table (2x2) to the document
table_2x2 = doc.add_table(rows=2, cols=2)



for index, row in dados.iterrows():
    linha_texto = "\t".join(map(str, row.values))
    print(linha_texto)

table_6x5 = doc.tables[0]
for row in range(6):
        table_6x5.cell(row, col).text = f"teste{row.values}"

        


# Acessar a primeira tabela (6x5) e preencher com um exemplo
table_6x5 = doc.tables[0]
for row in range(6):
    for col in range(5):
        table_6x5.cell(row, col).text = f"Linha {row+1}, Coluna {col+1}"

# Acessar a segunda tabela (2x2) e preencher com um exemplo
table_2x2 = doc.tables[1]
for row in range(2):
    for col in range(2):
        table_2x2.cell(row, col).text = f"Célula {row+1}, {col+1}"



# Save the document
output_path = "C:\\Users\\BurgerKing\\Desktop\\ProjetoA\\Document_with_Tables.docx"
doc.save(output_path)


